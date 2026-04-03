"""
DocuMind - AI-Powered Document Analysis API

A FastAPI application that extracts text from PDFs, DOCX, and images,
then analyzes them using Google Gemini for summarization, entity extraction,
and sentiment analysis.
"""

import base64
import io
import json
import os
import re

import fitz  # PyMuPDF
import pytesseract
from docx import Document
from dotenv import load_dotenv
from fastapi import FastAPI, File, Header, HTTPException, Request, UploadFile
from fastapi.responses import JSONResponse, RedirectResponse
from fastapi.staticfiles import StaticFiles
from google import genai
from google.genai import types
from PIL import Image
from pydantic import BaseModel


# ============================================================================
# CONFIGURATION
# ============================================================================

load_dotenv()

# Environment variables
API_KEY: str = os.getenv("API_KEY", "")
GEMINI_API_KEY: str = os.getenv("GEMINI_API_KEY", "")
ENVIRONMENT: str = os.getenv("ENVIRONMENT", "development")

# Validate required environment variables
if not API_KEY:
    raise RuntimeError("API_KEY is not found. Please set it in .env file")
if not GEMINI_API_KEY:
    raise RuntimeError("GEMINI_API_KEY is not found. Please set it in .env file")

# Initialize Google Gemini client
client = genai.Client(api_key=GEMINI_API_KEY)

# Model configuration with fallback
PRIMARY_MODEL = "gemini-3.0-flash"
FALLBACK_MODEL = "gemini-2.5-flash"

# Initialize FastAPI application
app = FastAPI(
    title="DocuMind API",
    description="AI-Powered Document Analysis API"
)

# Mount static files for web UI
app.mount("/static", StaticFiles(directory="static", html=True), name="static")


# ============================================================================
# DATA MODELS
# ============================================================================


class DocumentRequest(BaseModel):
    """Request model for document analysis endpoint."""
    fileName: str
    fileType: str  # pdf | docx | image
    fileBase64: str


# File extension to type mapping
EXTENSION_TO_TYPE = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".doc": "docx",
    ".png": "image",
    ".jpg": "image",
    ".jpeg": "image",
    ".tiff": "image",
    ".bmp": "image",
    ".webp": "image",
}


# ============================================================================
# AUTHENTICATION
# ============================================================================

def verify_api_key(x_api_key: str = Header(default=None)) -> None:
    """Verify API key from request header."""
    if not x_api_key or x_api_key != API_KEY:
        raise HTTPException(
            status_code=401,
            detail={
                "status": "error",
                "message": "Unauthorized: Invalid or missing API key.",
            },
        )


def is_local_request(request: Request) -> bool:
    """Check if the request originates from localhost."""
    client_host = request.client.host if request.client else ""
    return client_host in ("127.0.0.1", "localhost", "::1")


# ============================================================================
# TEXT EXTRACTION
# ============================================================================

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract text from PDF using PyMuPDF.
    Falls back to Tesseract OCR for scanned pages.
    
    Args:
        file_bytes: PDF file content as bytes
        
    Returns:
        Extracted text as string
    """
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    all_text: list[str] = []

    for page in doc:
        page_text = page.get_text("text").strip()

        if page_text:
            all_text.append(page_text)
        else:
            # Scanned page - render and apply OCR
            pix = page.get_pixmap(dpi=300)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            ocr_text = pytesseract.image_to_string(img, config="--psm 6").strip()
            if ocr_text:
                all_text.append(ocr_text)

    doc.close()
    return "\n\n".join(all_text)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract text from DOCX files using python-docx.
    Preserves reading order and includes table content.
    
    Args:
        file_bytes: DOCX file content as bytes
        
    Returns:
        Extracted text as string
    """
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                paragraphs.append(row_text)

    return "\n".join(paragraphs)


def extract_text_from_image(file_bytes: bytes) -> str:
    """
    Extract text from images using Tesseract OCR.
    
    Args:
        file_bytes: Image file content as bytes
        
    Returns:
        Extracted text as string
    """
    img = Image.open(io.BytesIO(file_bytes))
    return pytesseract.image_to_string(img, config="--psm 6").strip()


# ============================================================================
# AI ANALYSIS WITH GOOGLE GEMINI
# ============================================================================

ANALYSIS_PROMPT = """
You are a professional document analysis engine. Analyze the document text below and respond ONLY with a valid JSON object — no markdown, no backticks, no extra commentary.

The JSON must follow this exact schema:
{{
  "summary": "<concise 1-3 sentence summary of the document>",
  "entities": {{
    "names": ["<full person names found>"],
    "dates": ["<all dates found in any format>"],
    "organizations": ["<company, institution, or organisation names>"],
    "amounts": ["<monetary values with currency symbols>"]
  }},
  "sentiment": "<Positive | Negative | Neutral>"
}}

Rules:
- Extract ALL entities of each type present in the text following the rules below
- Do NOT return empty lists unless absolutely no entities exist in that category
- Summary must accurately reflect the document's purpose and key facts
- Do NOT include any text outside the JSON object
- Be precise and avoid false positives

Entity Extraction Rules:

Names:
- Extract full person names, authors, signatories, and individuals mentioned
- Examples: "John Smith", "Dr. Sarah Johnson", "Ravi Kumar"
- Ignore: Job titles alone (e.g., "Manager", "CEO"), pronouns, generic references

Dates:
- Only extract valid calendar dates (e.g., "June 2020", "2017-03-15", "March 2017", "10 March 2026")
- Ignore vague terms like "Present", "Current", "Ongoing", "To Date", "Recent", "Soon"
- Return dates in a consistent readable format: "Month YYYY" or "DD Month YYYY"

Organizations:
- Extract only real, specific organizations and meaningful organization groups actually named in the document
- Return only DISTINCT and MEANINGFUL organizations
- Include:
  * Companies (e.g., "ABC Pvt Ltd", "Google Inc", "Microsoft Corporation")
  * Institutions (e.g., "Harvard University", "Red Cross", "World Bank", "Reserve Bank")
  * Government bodies (e.g., "Ministry of Finance", "FBI", "Reserve Bank of India")
  * Organization groups (e.g., "financial institutions", "regulatory authorities", "healthcare providers")
  
- Do NOT include:
  * Tools/Software (e.g., "Python", "FastAPI", "Docker", "Tesseract", "PyMuPDF")
  * Frameworks/Libraries (e.g., "React", "TensorFlow", "pandas", "NumPy")
  * AI Models (e.g., "GPT-4", "Gemini", "Claude", "ChatGPT")
  * Programming languages (e.g., "JavaScript", "Python", "Java")
  * Technologies (e.g., "blockchain", "cloud computing", "machine learning")
  * Vague terms (e.g., "industry", "sector", "market", "field")
  * Overly generic or low-value terms (e.g., "companies", "technology companies", "private companies", "organizations")
  * Generic categories without context (e.g., "companies", "organizations" used generically)
  
- Rules for DISTINCT and MEANINGFUL extraction:
  * Remove duplicates and overlapping terms (e.g., if both "banks" and "financial institutions" appear, keep only the most informative version)
  * Prefer the most specific and informative version
  * Exclude overly generic standalone terms unless they are clearly used as meaningful entities in context
  
- Adapt extraction based on document type:
  * Resume/CV → extract specific company/institution names only (e.g., "Google", "Harvard University")
  * Reports/Articles → include meaningful organization groups (e.g., "regulatory authorities", "financial institutions")
  * Technical documentation → include only real-world organizations, ignore all tools/frameworks/technologies
  * Invoices/Receipts → extract company names mentioned (vendors, issuers)
  
- General rules:
  * Include meaningful real-world groups when they represent actual entities or sectors
  * Extract specific organizational entities actually mentioned in the document
  * Avoid duplicates and overly generic standalone terms
  * Prefer named entities over broad categories

Amounts:
- Extract monetary values with currency symbols (e.g., "₹10,000", "$500.50", "€1,200")
- Include complete amounts with commas/decimals as written
- Ignore: Percentages, counts, non-monetary numbers

Sentiment Rules:
- If the  document is factual (resume, CV, report, invoice, receipt, contract, form, technical documentation), return sentiment as "Neutral"
- Only return "Positive" or "Negative" for opinion-based content (reviews, feedback, complaint letters, testimonials with emotional tone)
- When in doubt, prefer "Neutral" for professional/business documents

Error Handling:
- If a field cannot be determined, use appropriate default: empty string for summary, empty arrays for entities, "Neutral" for sentiment
- Do not hallucinate or invent information not present in the document
- If the document text is unclear or corrupted, still attempt to extract whatever is readable
- Maintain consistent formatting across all extractions

Document text:
---
{text}
---
"""

def validate_gemini_response(data: dict) -> dict:
    """
    Validate and sanitize Gemini API response.
    Ensures response matches expected schema with proper defaults.
    
    Args:
        data: Raw response dictionary from Gemini
        
    Returns:
        Validated and sanitized response dictionary
    """
    validated = {
        "summary": "",
        "entities": {
            "names": [],
            "dates": [],
            "organizations": [],
            "amounts": [],
        },
        "sentiment": "Neutral",
    }
    
    # Extract summary
    if isinstance(data.get("summary"), str):
        validated["summary"] = data["summary"].strip()
    
    # Extract entities
    entities = data.get("entities", {})
    if isinstance(entities, dict):
        for key in ["names", "dates", "organizations", "amounts"]:
            if isinstance(entities.get(key), list):
                # Filter to only strings
                validated["entities"][key] = [
                    str(item).strip() for item in entities[key] 
                    if item and isinstance(item, (str, int, float))
                ]
    
    # Extract sentiment (must be one of the allowed values)
    sentiment = data.get("sentiment", "")
    if isinstance(sentiment, str) and sentiment.strip() in ("Positive", "Negative", "Neutral"):
        validated["sentiment"] = sentiment.strip()
    else:
        validated["sentiment"] = "Neutral"  # default fallback
    
    return validated


def analyse_with_gemini(text: str) -> dict:
    """
    Analyze document text using Google Gemini API.
    Implements automatic fallback from primary to secondary model on errors.
    
    Args:
        text: Document text to analyze (truncated to 12000 chars)
        
    Returns:
        Analysis results with summary, entities, and sentiment
        
    Raises:
        Exception: If all models fail to generate valid response
    """
    prompt = ANALYSIS_PROMPT.format(text=text[:12000])
    
    models_to_try = [PRIMARY_MODEL, FALLBACK_MODEL]
    last_error = None
    
    for model_name in models_to_try:
        try:
            response = client.models.generate_content(
                model=model_name,
                contents=prompt,
                config=types.GenerateContentConfig(
                    temperature=0.2,
                ),
            )
            raw = response.text.strip()
            
            # Clean response: remove markdown code fences if present
            raw = re.sub(r"^```(?:json)?", "", raw).strip()
            raw = re.sub(r"```$", "", raw).strip()
            
            return validate_gemini_response(json.loads(raw))
            
        except Exception as e:
            last_error = e
            error_str = str(e).lower()
            
            # Check if error is rate limit related
            if any(keyword in error_str for keyword in ["rate", "resource", "quota", "429"]):
                continue  # Try fallback model
                
            # For other errors, also attempt fallback
            continue
    
    # All models failed - raise the last error
    raise last_error


# ============================================================================
# API ENDPOINTS
# ============================================================================

@app.get("/")
def root():
    """Redirect root to web UI."""
    return RedirectResponse(url="/static/index.html")


@app.post("/api/document-analyze")
async def document_analyze(
    payload: DocumentRequest,
    x_api_key: str = Header(default=None),
):
    """
    Main API endpoint for document analysis.
    Accepts base64-encoded files and returns AI-powered analysis.
    
    Args:
        payload: DocumentRequest with fileName, fileType, and fileBase64
        x_api_key: API key for authentication (required)
        
    Returns:
        JSON response with summary, entities, and sentiment
        
    Raises:
        HTTPException: For authentication, validation, or processing errors
    """
    verify_api_key(x_api_key)

    # Validate fileType
    file_type = payload.fileType.lower().strip()
    supported = {"pdf", "docx", "image"}
    if file_type not in supported:
        raise HTTPException(
            status_code=400,
            detail={
                "status": "error",
                "message": f"Unsupported fileType '{payload.fileType}'. Use: pdf, docx, image.",
            },
        )

    # Decode base64
    try:
        file_bytes = base64.b64decode(payload.fileBase64)
    except Exception:
        raise HTTPException(
            status_code=400,
            detail={"status": "error", "message": "Invalid base64 encoding in fileBase64."},
        )

    # Extract text
    try:
        if file_type == "pdf":
            text = extract_text_from_pdf(file_bytes)
        elif file_type == "docx":
            text = extract_text_from_docx(file_bytes)
        else:
            text = extract_text_from_image(file_bytes)
    except Exception as e:
        raise HTTPException(
            status_code=422,
            detail={"status": "error", "message": f"Text extraction failed: {str(e)}"},
        )

    if not text.strip():
        raise HTTPException(
            status_code=422,
            detail={"status": "error", "message": "No readable text could be extracted from the document."},
        )

    # Gemini analysis
    try:
        analysis = analyse_with_gemini(text)
    except json.JSONDecodeError:
        raise HTTPException(
            status_code=502,
            detail={"status": "error", "message": "AI returned malformed response. Please retry."},
        )
    except Exception as e:
        raise HTTPException(
            status_code=502,
            detail={"status": "error", "message": f"AI analysis failed: {str(e)}"},
        )

    return JSONResponse(
        status_code=200,
        content={
            "status": "success",
            "fileName": payload.fileName,
            "summary": analysis.get("summary", ""),
            "entities": {
                "names": analysis.get("entities", {}).get("names", []),
                "dates": analysis.get("entities", {}).get("dates", []),
                "organizations": analysis.get("entities", {}).get("organizations", []),
                "amounts": analysis.get("entities", {}).get("amounts", []),
            },
            "sentiment": analysis.get("sentiment", "Neutral"),
        },
    )


@app.post("/api/upload-test")
async def upload_test(
    request: Request,
    file: UploadFile = File(...),
    x_api_key: str = Header(default=None),
):
    """
    Test endpoint for direct file upload.
    Accepts multipart/form-data files for easy testing.
    API key not required for localhost in development mode.
    
    Args:
        request: FastAPI request object
        file: Uploaded file
        x_api_key: API key (optional for localhost in dev mode)
        
    Returns:
        JSON response with analysis results
        
    Raises:
        HTTPException: For authentication, validation, or processing errors
    """
    # Check if running locally in development mode
    if ENVIRONMENT == "production" or not is_local_request(request):
        verify_api_key(x_api_key)

    # Detect file type from extension
    filename = file.filename or ""
    ext = os.path.splitext(filename)[-1].lower()
    file_type = EXTENSION_TO_TYPE.get(ext)

    if not file_type:
        raise HTTPException(
            status_code=400,
            detail={
                "status": "error",
                "message": f"Unsupported file extension '{ext}'. Supported: pdf, docx, png, jpg, jpeg, tiff, bmp, webp.",
            },
        )

    # Read file bytes
    file_bytes = await file.read()

    # Extract text
    try:
        if file_type == "pdf":
            text = extract_text_from_pdf(file_bytes)
        elif file_type == "docx":
            text = extract_text_from_docx(file_bytes)
        else:
            text = extract_text_from_image(file_bytes)
    except Exception as e:
        raise HTTPException(
            status_code=422,
            detail={"status": "error", "message": f"Text extraction failed: {str(e)}"},
        )

    if not text.strip():
        raise HTTPException(
            status_code=422,
            detail={"status": "error", "message": "No readable text could be extracted from the document."},
        )

    # Gemini analysis
    try:
        analysis = analyse_with_gemini(text)
    except json.JSONDecodeError:
        raise HTTPException(
            status_code=502,
            detail={"status": "error", "message": "AI returned malformed response. Please retry."},
        )
    except Exception as e:
        raise HTTPException(
            status_code=502,
            detail={"status": "error", "message": f"AI analysis failed: {str(e)}"},
        )

    return JSONResponse(
        status_code=200,
        content={
            "status": "success",
            "fileName": filename,
            "summary": analysis.get("summary", ""),
            "entities": {
                "names": analysis.get("entities", {}).get("names", []),
                "dates": analysis.get("entities", {}).get("dates", []),
                "organizations": analysis.get("entities", {}).get("organizations", []),
                "amounts": analysis.get("entities", {}).get("amounts", []),
            },
            "sentiment": analysis.get("sentiment", "Neutral"),
        },
    )